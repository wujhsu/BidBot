"""
文档加载器
Document loader utilities for processing PDF and DOCX files
"""

import os
import tempfile
import subprocess
from typing import List, Optional, Tuple
from pathlib import Path
import PyPDF2
import pypdf
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangchainDocument
from loguru import logger
import io

# 尝试导入pymupdf，如果失败则设置为None
try:
    import fitz  # pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    fitz = None
    PYMUPDF_AVAILABLE = False
    logger.warning("pymupdf未安装，将无法使用高级PDF处理功能")

class DocumentLoader:
    """文档加载器类"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        初始化文档加载器
        
        Args:
            chunk_size: 文本分块大小
            chunk_overlap: 文本分块重叠大小
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]
        )
    
    def load_pdf(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """
        加载PDF文件，使用多种PDF库的回退机制

        Args:
            file_path: PDF文件路径

        Returns:
            Tuple[str, List[Tuple[str, int]]]: (全文内容, [(页面内容, 页码)])
        """
        # 尝试多种PDF处理方法
        methods = [
            ("pypdf", self._load_pdf_with_pypdf),
            ("PyPDF2", self._load_pdf_with_pypdf2),
            ("修复模式", self._load_pdf_with_repair)
        ]

        # 如果pymupdf可用，添加到方法列表
        if PYMUPDF_AVAILABLE:
            methods.insert(0, ("pymupdf", self._load_pdf_with_pymupdf))

        last_error = None
        for method_name, method_func in methods:
            try:
                logger.info(f"尝试使用{method_name}加载PDF: {file_path}")
                full_text, pages_content = method_func(file_path)

                if full_text.strip():  # 确保提取到了内容
                    logger.info(f"使用{method_name}成功加载PDF文件: {file_path}, 共{len(pages_content)}页")
                    return full_text, pages_content
                else:
                    logger.warning(f"{method_name}提取的内容为空，尝试下一种方法")

            except Exception as e:
                logger.warning(f"{method_name}加载失败: {e}")
                last_error = e
                continue

        # 所有方法都失败了
        error_msg = f"所有PDF加载方法都失败了，最后一个错误: {last_error}"
        logger.error(error_msg)
        raise Exception(error_msg)

    def _load_pdf_with_pypdf(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """使用pypdf库加载PDF"""
        full_text = ""
        pages_content = []

        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        full_text += f"\n\n--- 第{page_num}页 ---\n\n{page_text}"
                        pages_content.append((page_text, page_num))
                    else:
                        logger.warning(f"第{page_num}页无法提取文本，可能是扫描件")
                        pages_content.append(("", page_num))
                except Exception as e:
                    logger.error(f"处理第{page_num}页时出错: {e}")
                    pages_content.append(("", page_num))

        return full_text, pages_content

    def _load_pdf_with_pypdf2(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """使用PyPDF2库加载PDF"""
        full_text = ""
        pages_content = []

        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)

            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        full_text += f"\n\n--- 第{page_num}页 ---\n\n{page_text}"
                        pages_content.append((page_text, page_num))
                    else:
                        logger.warning(f"第{page_num}页无法提取文本，可能是扫描件")
                        pages_content.append(("", page_num))
                except Exception as e:
                    logger.error(f"处理第{page_num}页时出错: {e}")
                    pages_content.append(("", page_num))

        return full_text, pages_content

    def _load_pdf_with_repair(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """使用修复模式加载PDF"""
        full_text = ""
        pages_content = []

        try:
            # 尝试修复PDF文件
            with open(file_path, 'rb') as file:
                pdf_data = file.read()

            # 使用io.BytesIO创建内存中的PDF对象
            pdf_stream = io.BytesIO(pdf_data)

            # 尝试使用pypdf的严格模式关闭
            pdf_reader = pypdf.PdfReader(pdf_stream, strict=False)

            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        full_text += f"\n\n--- 第{page_num}页 ---\n\n{page_text}"
                        pages_content.append((page_text, page_num))
                    else:
                        logger.warning(f"第{page_num}页无法提取文本，可能是扫描件")
                        pages_content.append(("", page_num))
                except Exception as e:
                    logger.error(f"处理第{page_num}页时出错: {e}")
                    pages_content.append(("", page_num))

            return full_text, pages_content

        except Exception as e:
            logger.error(f"修复模式加载失败: {e}")
            raise

    def _load_pdf_with_pymupdf(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """使用pymupdf库加载PDF（最强大的PDF处理库）"""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("pymupdf未安装")

        full_text = ""
        pages_content = []

        try:
            # 打开PDF文档
            doc = fitz.open(file_path)

            for page_num in range(len(doc)):
                try:
                    page = doc.load_page(page_num)
                    page_text = page.get_text()

                    if page_text.strip():
                        full_text += f"\n\n--- 第{page_num + 1}页 ---\n\n{page_text}"
                        pages_content.append((page_text, page_num + 1))
                    else:
                        logger.warning(f"第{page_num + 1}页无法提取文本，可能是扫描件")
                        pages_content.append(("", page_num + 1))

                except Exception as e:
                    logger.error(f"处理第{page_num + 1}页时出错: {e}")
                    pages_content.append(("", page_num + 1))

            doc.close()
            return full_text, pages_content

        except Exception as e:
            logger.error(f"pymupdf加载失败: {e}")
            raise
    
    def _convert_docx_to_pdf_temp(self, docx_path: str) -> Optional[str]:
        """
        将DOCX文件临时转换为PDF以获取真实页码

        Args:
            docx_path: DOCX文件路径

        Returns:
            Optional[str]: 临时PDF文件路径，如果转换失败则返回None
        """
        try:
            # 创建临时PDF文件
            temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
            temp_pdf_path = temp_pdf.name
            temp_pdf.close()

            # 尝试使用LibreOffice转换（如果可用）
            try:
                result = subprocess.run([
                    'libreoffice', '--headless', '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(temp_pdf_path),
                    docx_path
                ], capture_output=True, text=True, timeout=30)

                if result.returncode == 0:
                    # LibreOffice会生成与原文件同名的PDF
                    docx_name = Path(docx_path).stem
                    generated_pdf = os.path.join(os.path.dirname(temp_pdf_path), f"{docx_name}.pdf")

                    if os.path.exists(generated_pdf):
                        # 重命名为我们的临时文件名
                        os.rename(generated_pdf, temp_pdf_path)
                        logger.info(f"使用LibreOffice成功转换DOCX为PDF: {temp_pdf_path}")
                        return temp_pdf_path

            except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.SubprocessError) as e:
                logger.warning(f"LibreOffice转换失败: {e}")

            # 如果LibreOffice不可用，尝试使用python-docx2pdf（如果安装了）
            try:
                from docx2pdf import convert
                convert(docx_path, temp_pdf_path)
                if os.path.exists(temp_pdf_path) and os.path.getsize(temp_pdf_path) > 0:
                    logger.info(f"使用docx2pdf成功转换DOCX为PDF: {temp_pdf_path}")
                    return temp_pdf_path
            except ImportError:
                logger.warning("docx2pdf未安装，无法使用此方法转换")
            except Exception as e:
                logger.warning(f"docx2pdf转换失败: {e}")

            # 清理失败的临时文件
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)

            return None

        except Exception as e:
            logger.error(f"DOCX转PDF过程中发生错误: {e}")
            return None

    def _get_docx_paragraph_to_page_mapping(self, docx_path: str) -> dict:
        """
        获取DOCX文件中段落到页码的映射

        Args:
            docx_path: DOCX文件路径

        Returns:
            dict: 段落号到页码的映射 {段落号: 页码}
        """
        paragraph_to_page = {}

        # 尝试转换为PDF获取真实页码
        temp_pdf_path = self._convert_docx_to_pdf_temp(docx_path)

        if temp_pdf_path and PYMUPDF_AVAILABLE:
            try:
                # 读取原始DOCX文件的段落
                doc = Document(docx_path)
                docx_paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        docx_paragraphs.append(text)

                # 读取转换后的PDF文件
                pdf_doc = fitz.open(temp_pdf_path)

                # 为每个段落找到对应的页码
                for para_num, para_text in enumerate(docx_paragraphs, 1):
                    # 在PDF中搜索这个段落的文本
                    found_page = None
                    search_text = para_text[:50]  # 使用前50个字符进行搜索

                    for page_num in range(len(pdf_doc)):
                        page = pdf_doc.load_page(page_num)
                        text_instances = page.search_for(search_text)
                        if text_instances:
                            found_page = page_num + 1
                            break

                    if found_page:
                        paragraph_to_page[para_num] = found_page
                    else:
                        # 如果找不到，使用估算方法
                        estimated_page = max(1, (para_num - 1) // 25 + 1)
                        paragraph_to_page[para_num] = estimated_page

                pdf_doc.close()
                logger.info(f"成功建立段落到页码映射，共{len(paragraph_to_page)}个段落")

            except Exception as e:
                logger.error(f"建立段落页码映射失败: {e}")
            finally:
                # 清理临时PDF文件
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)

        # 如果无法获取真实页码，使用估算方法
        if not paragraph_to_page:
            logger.warning("无法获取真实页码，使用估算方法")
            doc = Document(docx_path)
            para_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    para_count += 1
                    estimated_page = max(1, (para_count - 1) // 25 + 1)
                    paragraph_to_page[para_count] = estimated_page

        return paragraph_to_page

    def load_docx(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """
        加载DOCX文件，尝试获取真实页码

        Args:
            file_path: DOCX文件路径

        Returns:
            Tuple[str, List[Tuple[str, int]]]: (全文内容, [(段落内容, 页码)])
        """
        try:
            # 获取段落到页码的映射
            paragraph_to_page = self._get_docx_paragraph_to_page_mapping(file_path)

            doc = Document(file_path)
            full_text = ""
            paragraphs_content = []

            para_count = 0
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    para_count += 1
                    page_number = paragraph_to_page.get(para_count, 1)

                    # 在文本中标记页码而不是段落号
                    full_text += f"\n\n--- 第{page_number}页 ---\n\n{para_text}"
                    paragraphs_content.append((para_text, page_number))

            logger.info(f"成功加载DOCX文件: {file_path}, 共{len(paragraphs_content)}段，映射到{len(set(paragraph_to_page.values()))}页")
            return full_text, paragraphs_content

        except Exception as e:
            logger.error(f"加载DOCX文件失败: {file_path}, 错误: {e}")
            # 如果新方法失败，回退到原始方法
            return self._load_docx_fallback(file_path)

    def _load_docx_fallback(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """
        DOCX文件加载的回退方法（使用段落号估算页码）

        Args:
            file_path: DOCX文件路径

        Returns:
            Tuple[str, List[Tuple[str, int]]]: (全文内容, [(段落内容, 估算页码)])
        """
        try:
            doc = Document(file_path)
            full_text = ""
            paragraphs_content = []

            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text.strip()
                if para_text:
                    # 使用估算页码
                    estimated_page = max(1, (para_num - 1) // 25 + 1)
                    full_text += f"\n\n--- 第{estimated_page}页 ---\n\n{para_text}"
                    paragraphs_content.append((para_text, estimated_page))

            logger.info(f"使用回退方法加载DOCX文件: {file_path}, 共{len(paragraphs_content)}段")
            return full_text, paragraphs_content

        except Exception as e:
            logger.error(f"回退方法加载DOCX文件失败: {file_path}, 错误: {e}")
            raise
    
    def load_txt(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """
        加载TXT文件

        Args:
            file_path: TXT文件路径

        Returns:
            Tuple[str, List[Tuple[str, int]]]: (全文内容, [(行内容, 行号)])
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()

            full_text = ""
            lines_content = []

            for line_num, line in enumerate(lines, 1):
                line_text = line.strip()
                if line_text:
                    full_text += f"\n\n--- 第{line_num}行 ---\n\n{line_text}"
                    lines_content.append((line_text, line_num))

            logger.info(f"成功加载TXT文件: {file_path}, 共{len(lines_content)}行")
            return full_text, lines_content

        except Exception as e:
            logger.error(f"加载TXT文件失败: {file_path}, 错误: {e}")
            raise

    def load_document(self, file_path: str) -> Tuple[str, List[Tuple[str, int]]]:
        """
        根据文件扩展名自动选择加载方法

        Args:
            file_path: 文件路径

        Returns:
            Tuple[str, List[Tuple[str, int]]]: (全文内容, [(内容, 位置号)])
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_extension = Path(file_path).suffix.lower()

        if file_extension == '.pdf':
            return self.load_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.load_docx(file_path)
        elif file_extension == '.txt':
            return self.load_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
    
    def split_text(self, text: str, metadata: Optional[dict] = None) -> List[LangchainDocument]:
        """
        将文本分割成块

        Args:
            text: 要分割的文本
            metadata: 元数据

        Returns:
            List[LangchainDocument]: 分割后的文档块列表
        """
        if metadata is None:
            metadata = {}

        # 使用文本分割器分割文本
        chunks = self.text_splitter.split_text(text)

        # 创建LangchainDocument对象
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.copy()
            doc_metadata.update({
                "chunk_id": i,
                "chunk_size": len(chunk)
            })

            # 提取页码/段落号信息并添加到元数据
            page_info = self._extract_location_info(chunk)
            if page_info:
                doc_metadata.update(page_info)

            documents.append(LangchainDocument(page_content=chunk, metadata=doc_metadata))

        logger.info(f"文本分割完成，共生成{len(documents)}个文档块")
        return documents

    def _extract_location_info(self, text: str) -> dict:
        """
        从文本块中提取位置信息（页码或段落号）

        Args:
            text: 文本块

        Returns:
            dict: 包含位置信息的字典
        """
        import re
        location_info = {}

        # 查找页码标记（PDF文件和改进后的DOCX文件）
        page_pattern = r'--- 第(\d+)页 ---'
        page_matches = re.findall(page_pattern, text)
        if page_matches:
            # 如果有多个页码，取第一个
            location_info['page_number'] = int(page_matches[0])
            location_info['location_type'] = 'page'

        # 查找段落标记（旧版DOCX文件处理方式，作为回退）
        para_pattern = r'--- 第(\d+)段 ---'
        para_matches = re.findall(para_pattern, text)
        if para_matches and 'page_number' not in location_info:
            # 如果有多个段落号，取第一个
            para_num = int(para_matches[0])
            # 将段落号转换为估算页码
            estimated_page = max(1, (para_num - 1) // 25 + 1)
            location_info['page_number'] = estimated_page
            location_info['paragraph_number'] = para_num
            location_info['location_type'] = 'paragraph'

        # 查找行号标记（TXT文件）
        line_pattern = r'--- 第(\d+)行 ---'
        line_matches = re.findall(line_pattern, text)
        if line_matches and 'page_number' not in location_info:
            line_num = int(line_matches[0])
            # 将行号转换为估算页码
            estimated_page = max(1, (line_num - 1) // 50 + 1)
            location_info['page_number'] = estimated_page
            location_info['line_number'] = line_num
            location_info['location_type'] = 'line'

        return location_info
    
    def process_document(self, file_path: str) -> Tuple[str, List[LangchainDocument]]:
        """
        处理文档：加载并分割
        
        Args:
            file_path: 文件路径
            
        Returns:
            Tuple[str, List[LangchainDocument]]: (全文内容, 文档块列表)
        """
        # 加载文档
        full_text, content_parts = self.load_document(file_path)
        
        # 准备元数据
        metadata = {
            "source": file_path,
            "file_name": Path(file_path).name,
            "file_type": Path(file_path).suffix.lower()
        }
        
        # 分割文本
        documents = self.split_text(full_text, metadata)
        
        return full_text, documents
